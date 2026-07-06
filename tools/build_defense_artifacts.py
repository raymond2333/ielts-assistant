from __future__ import annotations

import copy
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PPTX_PATH = Path(
    "/Users/raymondtse/Downloads/学习科研赛道-深圳校区-信达雅-基于多智能体协作的智能雅思学习助手-谢镓骏/信达雅答辩版.pptx"
)
OUT_DIR = Path("/Users/raymondtse/Desktop/code/python/llm/ielts/outputs")
MODIFIED_PPTX = OUT_DIR / "信达雅答辩版_备注修改稿.pptx"
DOCX_PATH = OUT_DIR / "信达雅答辩训练稿.docx"

NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
}

for prefix, uri in NS.items():
    if prefix not in {"rel", "ct"}:
        ET.register_namespace(prefix, uri)


SLIDE_TITLES = {
    1: "开场：信达雅 IELTS Assistant",
    2: "目录",
    3: "项目背景：市场痛点",
    4: "范式升级",
    5: "解决方案总览",
    6: "产品功能",
    7: "口语训练",
    8: "串题功能",
    9: "写作训练",
    10: "学习规划",
    11: "多智能体架构总览",
    12: "Agent 协作模式",
    13: "结构化 Prompt 工程",
    14: "学习记忆闭环",
    15: "七天学习全景",
    16: "竞争优势",
    17: "未来规划",
    18: "团队分工与交付",
    19: "比赛完成度总结",
    20: "结束页",
}


NOTES = {
    1: "各位评委老师好，我是深圳校区信达雅团队的队长谢镓骏。今天汇报的项目是“信达雅”，一个基于多智能体协作的智能雅思学习助手。我们的核心判断是：雅思备考缺的不是又一个聊天窗口，而是一个能分工、能记忆、能持续陪伴的 AI 教学团队。所以信达雅把出题、教练、评测、表达、记忆和规划拆成六类智能体，由系统统一调度，覆盖从练习、反馈到复盘的完整链路。接下来我会结合我们已经实现的工程系统，说明它为什么不是概念 Demo，而是一个可以部署、可以使用、可以积累学习数据的产品。",
    2: "我的汇报会按照五个部分展开。第一部分讲项目背景，说明我们为什么选择雅思备考这个场景；第二部分讲解决方案和产品功能，展示信达雅如何把多 Agent 能力落到具体学习流程里；第三部分讲核心技术，包括智能体架构、结构化 Prompt 和记忆闭环；第四部分讲应用落实，也就是七天训练、竞争优势和未来路线；最后介绍团队分工与交付完成度。整体逻辑是从问题出发，回到工程实现，再落到可验证的产品成果。",
    3: "先看市场痛点。雅思考生真正缺的不是资料，而是反馈、规划和数据三个断点。反馈断点是口语练完没人指出问题，写作批改也往往要等一到三天；规划断点是学生知道要努力，却不知道下一步该练什么；数据断点是大量练习做完就结束，没有沉淀成学习档案。信达雅针对这三个断点分别给出即时四维评分、个性化周计划和 MySQL 持久化记录。也就是说，我们不是把 AI 当问答工具，而是把 AI 组织成一个持续陪伴学生的训练系统。",
    4: "这一页想表达的是备考范式的变化。传统方式里，学生要在 ChatGPT、词典、翻译、批改网站和背词软件之间手动搬运信息，练习数据被切碎在不同工具里。信达雅的思路是把这些任务收束到一个统一系统中：前端负责顺手的学习体验，后端由智能体处理出题、批改、串题和规划，底层用学习记忆库记录用户画像、训练结果和词汇资产。这样学生面对的不再是工具堆叠，而是一个能自动衔接学习流程的 AI 教学团队。",
    5: "信达雅的整体方案可以概括为“多智能体协同编排加学习记忆闭环”。这里有五项核心创新。第一是多 Agent 协作调度，把雅思任务拆给专业角色；第二是 14 个结构化 Prompt 模板，让模型输出能被前端解析和下游使用；第三是学习记忆数据飞轮，每次练习都保存为结构化资产；第四是口语串题引擎，把多个 Part 2 话题压缩成可迁移的核心故事；第五是产品化落地。我们已经实现一键启动脚本、双前端、多模型配置和 MySQL 持久化，所以它不是只跑一次的课堂 Demo。",
    6: "产品层面，我们刻意没有做成一个“请输入提示词”的聊天框，而是做成学生可以直接使用的学习工具。学习总览展示当前水平、目标分数和最近训练；口语训练支持 Part 1、Part 2、Part 3，包含 TTS 朗读、计时、文本和语音输入；写作批改支持 Task 1 和 Task 2；串题中心生成可迁移故事线；词汇模块支持查词、收藏和复习；后台则可以管理用户和训练记录。学生看到的是顺畅的界面，背后是多个 Agent 在自动完成任务分工。",
    7: "口语训练是最能体现多 Agent 协作的场景。学生选择题型、话题和难度后，出题 Agent 先生成题目和参考表达；如果学生没有思路，教练 Agent 可以把中文思路转成英文口语答案；学生提交答案后，评测 Agent 按雅思四项标准给出总分、分项分析和优化回答；最后记忆 Agent 把题目、作答、评分和反馈保存下来，规划 Agent 后续会读取这些记录调整建议。也就是说，学生只完成一次作答，系统已经完成了出题、辅导、诊断和复盘的完整闭环。",
    8: "串题功能是我们差异化最大的功能。雅思 Part 2 题库很大，但很多题其实可以共用一个核心经历。传统方法是逐题背诵，五道题可能要准备一千五百词；信达雅的串题 Agent 会提取统一核心故事，再为每个题目生成不同侧重点的适配版本，记忆量可以明显下降。它的价值不只是生成答案，而是把学生的表达变成可复用、可迁移的资产。",
    9: "写作训练展示的是评测体系的专业化。学生可以先让出题 Agent 生成 Task 1 或 Task 2 题目，其中 Task 1 支持柱状图、线形图、饼图、流程图、表格和地图等类型。提交作文后，评测 Agent 按官方四项标准评分：任务完成度、连贯与衔接、词汇资源、语法准确性，并给出逐项修改建议。表达 Agent 还可以生成同题参考范文，帮助学生把“哪里错了”和“怎么写更好”连起来。",
    10: "学习规划解决的是“下一步练什么”的问题。传统课表是固定的，学生练了很多题，却不知道弱项是否真的被修复。信达雅会综合用户画像、目标分数、考试日期和历史训练记录，生成个性化计划。这里的规划不是凭空建议，而是建立在真实训练数据上的：最近哪些题型分数波动大、哪个维度持续偏弱、距离目标分还有多远，都会影响下一步推荐。",
    11: "这一页是系统架构。最上层是用户交互层，我们同时保留新版 Web 入口和稳定版入口；第二层是智能编排层，负责把用户动作路由到相应智能体能力；第三层是专业 Agent 层，六类任务分别对应出题、教练、评测、表达、记忆和规划；最下层是 MySQL 持久记忆层，负责保存用户画像、训练记录、单词掌握状态和生词本。这个结构让信达雅既有 AI 能力，也有产品系统的稳定性。",
    12: "这里要强调，Agent 协作不是让几个模型互相聊天，而是结构化任务分解、有向数据流和统一记忆。第一种模式是串行链式，比如口语从出题到评分再到保存；第二种是并行分发，比如写作批改同时关注任务、连贯、词汇和语法四个维度，再汇总成报告；第三种是事件触发，比如保存训练记录后，学习分析页可以刷新重点建议和计划。系统内部传递的是题目、作答、分数、弱项等结构化信息，而不是完全不可控的自由聊天文本。",
    13: "结构化 Prompt 是系统稳定性的关键。我们不是简单写一句“请帮我批改作文”，而是为不同 Agent 定义角色、任务、输出结构和约束。比如口语出题要稳定给出题目、参考答案、关键词和技巧；口语反馈要稳定给出总分、四项分析、优化回答和练习建议；写作批改也要对齐 IELTS 官方四维评分。这样前端可以稳定渲染评分卡、修改建议、范文和词汇推荐，下游 Agent 也能继续消费这些结果。",
    14: "学习记忆闭环是信达雅区别于普通聊天工具的核心。普通聊天工具最多保存一段对话，而我们保存的是学习资产。系统会记录用户画像、每次训练结果、单词掌握状态和生词本。每次口语反馈、写作批改、串题方案或学习计划都会沉淀为结构化学习数据。下一次生成计划时，系统读取历史记录，判断最近分数、弱项和练习频率，让建议越来越贴近个人。",
    15: "把系统放到七天使用场景里会更直观。第一天，学生录入当前水平和目标分数，规划 Agent 生成阶段计划；第二天做 Part 2 口语，出题、教练、评测和记忆 Agent 完成一次闭环；第三天根据反馈复练，系统记录二次评分；第四天从反馈里提取生词并进入词汇复习；第五到六天切换到 Task 1、Task 2 和串题训练；第七天，规划 Agent 读取一周记录，更新下一周重点。这个过程说明信达雅不是一次性回答，而是在持续训练学生。",
    16: "竞争优势可以概括为三点。第一，多 Agent 协作让不同任务专业化，而不是一个通用模型包办所有事；第二，雅思垂直 Prompt 工程把官方评分标准、题型约束和输出结构固化下来；第三，数据飞轮让每次练习都进入长期学习档案。通用 AI 可以回答问题，但通常没有稳定的题型流程、历史记录和弱项规划；传统培训有老师，但成本高、反馈频次低。信达雅的定位是在低成本下提供高频、结构化、可追踪的训练体验。",
    17: "未来规划上，我们认为这套架构不只适用于雅思，而是可以迁移到更多语言学习场景。短期会补齐发音评分、听力阅读专项和题库校准；中期会做移动端、小程序、班级学情和教师端；长期可以扩展到托福、四六级和考研英语。原因是底层能力是通用的：出题、训练、评测、记忆、规划这些环节在多数语言考试里都存在。只要替换垂直 Prompt、题库标准和评分维度，多 Agent 架构就能复用。",
    18: "项目能落地，背后是明确的团队分工。后端架构负责整体系统、LangChain 调用、多模型适配和数据库；前端全栈负责 Flask、Streamlit 两套界面和交互组件；Prompt 工程负责 14 套模板、评分标准对齐和输出约束；测试部署负责功能测试、一键启动脚本、文档和演示材料。我们不是只交了一份 PPT，而是交付了可运行代码、数据库 Schema、前端页面、演示视频和部署文档。答辩现场如果需要，也可以从代码和页面两侧验证完成度。",
    19: "最后总结一下。信达雅的创意性体现在把“一个 AI 做所有事”改造成“多个专业 Agent 协作”；实用性体现在口语、写作、串题、词汇、规划都落到了真实操作流程；完成度体现在 Flask 和 Streamlit 双版本、MySQL 持久化、多模型热切换、HMAC 跨版本免登录和一键部署；技术合理性体现在结构化 Prompt、雅思四维评分、数据契约和学习记忆闭环。我们的结论是：信达雅不是通用 AI 加雅思提示词，而是围绕雅思备考重新设计的 Agent 工作流系统。",
    20: "以上就是信达雅团队的汇报。我们希望用多智能体协作，把雅思备考从碎片化工具堆叠，升级为一个有分工、有记忆、有规划的 AI 教学团队。感谢各位评委老师的聆听，欢迎老师们提问，也欢迎现场扫码查看源代码、体验 Demo 或观看演示视频。",
}


def set_paragraph_text(shape, text: str) -> None:
    tx_body = shape.find("p:txBody", NS)
    if tx_body is None:
        tx_body = ET.SubElement(shape, f"{{{NS['p']}}}txBody")
        ET.SubElement(tx_body, f"{{{NS['a']}}}bodyPr")
        ET.SubElement(tx_body, f"{{{NS['a']}}}lstStyle")
    for child in list(tx_body):
        if child.tag == f"{{{NS['a']}}}p":
            tx_body.remove(child)
    para = ET.SubElement(tx_body, f"{{{NS['a']}}}p")
    run = ET.SubElement(para, f"{{{NS['a']}}}r")
    r_pr = ET.SubElement(run, f"{{{NS['a']}}}rPr")
    r_pr.set("lang", "zh-CN")
    t = ET.SubElement(run, f"{{{NS['a']}}}t")
    t.text = text


def update_notes_xml(xml_bytes: bytes, text: str) -> bytes:
    root = ET.fromstring(xml_bytes)
    body_shape = None
    for shape in root.findall(".//p:sp", NS):
        ph = shape.find(".//p:ph", NS)
        if ph is not None and ph.get("type") == "body":
            body_shape = shape
            break
    if body_shape is None:
        body_shape = root.find(".//p:sp", NS)
    set_paragraph_text(body_shape, text)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def make_notes_slide20(template_xml: bytes, text: str) -> bytes:
    return update_notes_xml(template_xml, text)


def add_notes_relationship(rels_xml: bytes) -> bytes:
    root = ET.fromstring(rels_xml)
    for rel in root.findall("rel:Relationship", NS):
        if rel.get("Type", "").endswith("/notesSlide"):
            rel.set("Target", "../notesSlides/notesSlide20.xml")
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    used_ids = {rel.get("Id") for rel in root.findall("rel:Relationship", NS)}
    idx = 1
    while f"rId{idx}" in used_ids:
        idx += 1
    rel = ET.SubElement(root, f"{{{NS['rel']}}}Relationship")
    rel.set("Id", f"rId{idx}")
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/notesSlide")
    rel.set("Target", "../notesSlides/notesSlide20.xml")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def add_content_type(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    part_name = "/ppt/notesSlides/notesSlide20.xml"
    for override in root.findall("ct:Override", NS):
        if override.get("PartName") == part_name:
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    override = ET.SubElement(root, f"{{{NS['ct']}}}Override")
    override.set("PartName", part_name)
    override.set("ContentType", "application/vnd.openxmlformats-officedocument.presentationml.notesSlide+xml")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def build_pptx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(PPTX_PATH, MODIFIED_PPTX)
    replacements = {}
    with zipfile.ZipFile(MODIFIED_PPTX, "r") as zin:
        for i, note in NOTES.items():
            name = f"ppt/notesSlides/notesSlide{i}.xml"
            if name in zin.namelist():
                replacements[name] = update_notes_xml(zin.read(name), note)
        replacements["ppt/notesSlides/notesSlide20.xml"] = make_notes_slide20(
            zin.read("ppt/notesSlides/notesSlide19.xml"), NOTES[20]
        )
        replacements["ppt/notesSlides/_rels/notesSlide20.xml.rels"] = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            b'<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="../slides/slide20.xml"/>'
            b'<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/notesMaster" Target="../notesMasters/notesMaster1.xml"/>'
            b"</Relationships>"
        )
        replacements["ppt/slides/_rels/slide20.xml.rels"] = add_notes_relationship(
            zin.read("ppt/slides/_rels/slide20.xml.rels")
        )
        replacements["[Content_Types].xml"] = add_content_type(zin.read("[Content_Types].xml"))

        tmp = MODIFIED_PPTX.with_suffix(".tmp.pptx")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            seen = set()
            for item in zin.infolist():
                data = replacements.get(item.filename, zin.read(item.filename))
                zout.writestr(item, data)
                seen.add(item.filename)
            for name, data in replacements.items():
                if name not in seen:
                    zout.writestr(name, data)
    tmp.replace(MODIFIED_PPTX)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)


def set_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 14)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx() -> None:
    doc = Document()
    set_doc_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("信达雅答辩训练稿")
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph("基于工程实例重写的 PPT 备注稿；仅调整讲稿，不改动可见幻灯片内容。")
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("训练方式", level=1)
    for item in [
        "先按页顺序通读一遍，确认每页只讲一个核心观点。",
        "正式练习时控制总时长在 8 到 10 分钟，技术页可以略放慢。",
        "遇到评委追问时，优先回到可验证能力：多模型接入、结构化评分、学习记忆库、双前端和一键部署。",
    ]:
        doc.add_paragraph(item, style=None).style = doc.styles["Normal"]

    doc.add_heading("逐页讲稿", level=1)
    for i in range(1, 21):
        doc.add_heading(f"第 {i} 页：{SLIDE_TITLES[i]}", level=2)
        p = doc.add_paragraph(NOTES[i])
        p.paragraph_format.first_line_indent = Pt(22)
        if i in {5, 10, 11, 14, 18, 19}:
            tip = {
                5: "提示：这一页讲“五项创新”，注意口径统一为页面上的“14 个结构化 Prompt”。",
                10: "提示：原备注只有“改”，这一页现在重点讲真实的规划实现和兜底逻辑。",
                11: "提示：被追问架构时，讲四层结构即可：交互层、编排层、智能体层、记忆层。",
                14: "提示：把“聊天记录”和“结构化学习数据”区分开，这是回答产品壁垒的关键。",
                18: "提示：团队页不要只报人名，要突出每个人有可验证交付。",
                19: "提示：收束到“Agent 工作流系统”，避免听起来像普通 AI 套壳。",
            }[i]
            note = doc.add_paragraph()
            note_run = note.add_run(tip)
            note_run.bold = True
            note_run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_page_break()
    doc.add_heading("答辩追问速查", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["答辩点", "可验证能力", "讲法"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "F2F4F7")
    rows = [
        ("多模型适配", "多供应商与兼容接口", "支持通义、DeepSeek、OpenAI 和兼容接口，不绑定单一模型。"),
        ("口语全流程", "出题、辅助作答、四维评分、记录沉淀", "学生只完成一次作答，系统自动形成训练闭环。"),
        ("写作批改", "Task 1/2 出题、批改、范文对比", "按 IELTS 四维标准输出评分、建议和参考范文。"),
        ("串题引擎", "核心故事线提取与多题适配", "把多个 Part 2 话题提炼为可迁移表达资产。"),
        ("学习记忆", "用户画像、训练记录、词汇资产", "保存的是结构化学习数据，不是普通聊天历史。"),
        ("产品完成度", "双版本入口、持久化、一键启动", "可部署、可登录、可保存数据、可现场演示。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_pptx()
    build_docx()
    print(MODIFIED_PPTX)
    print(DOCX_PATH)
